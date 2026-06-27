"""
parsing.py
----------
Turns an uploaded resume file (PDF / DOCX / TXT) into plain text.

Each format needs a different library:
  - PDF  -> pdfplumber (reliable text extraction, handles layout decently)
  - DOCX -> python-docx (reads the XML inside the .docx zip)
  - TXT  -> just decode the bytes
"""

import io
import pdfplumber
from docx import Document


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Dispatch to the right extractor based on file extension."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _from_pdf(file_bytes)
    if name.endswith(".docx"):
        return _from_docx(file_bytes)
    if name.endswith((".txt", ".text")):
        return file_bytes.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {filename}")


def _from_pdf(file_bytes: bytes) -> str:
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs]
    # Also pull text out of tables, which python-docx keeps separate.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)
